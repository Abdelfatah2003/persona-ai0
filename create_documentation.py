from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Configure page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Helper functions
def add_heading(text, level=1, center=False):
    heading = doc.add_heading(text, level=level)
    if center:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(99, 102, 241)
    return heading

def add_paragraph(text='', bold=False, indent=0):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_table_with_headers(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
    
    return table

# ==========================================
# COVER PAGE
# ==========================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Personality & Career Recommendation System')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(99, 102, 241)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Smart System for Analyzing User Personality and\nRecommending Career Paths and Friends')
run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

prepared = doc.add_paragraph()
prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prepared.add_run('Prepared for Graduation Project Submission')
run.font.size = Pt(14)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Student info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('[Student Name]\n[Student ID]\n[Department of Computer Science]\n[University Name]\n[2026]')
run.font.size = Pt(12)

doc.add_page_break()

# ==========================================
# TABLE OF CONTENTS
# ==========================================
add_heading('Table of Contents', 1, center=True)
doc.add_paragraph()

toc_items = [
    ('1.', 'Project Introduction', ''),
    ('1.1', 'Project Overview', ''),
    ('1.2', 'Project Name', ''),
    ('1.3', 'Project Objectives', ''),
    ('1.4', 'Project Scope', ''),
    ('2.', 'System Analysis', ''),
    ('2.1', 'Requirements Analysis', ''),
    ('2.1.1', 'Functional Requirements', ''),
    ('2.1.2', 'Non-Functional Requirements', ''),
    ('2.2', 'System Context Diagram', ''),
    ('2.3', 'Use Case Diagram', ''),
    ('2.4', 'Use Case Descriptions', ''),
    ('2.5', 'Wireframes', ''),
    ('3.', 'System Design', ''),
    ('3.1', 'Database Design', ''),
    ('3.2', 'Entity Relationship Diagram (ERD)', ''),
    ('3.3', 'Database Mapping', ''),
    ('3.4', 'Service Architecture', ''),
    ('4.', 'Implementation', ''),
    ('4.1', 'Frontend Implementation', ''),
    ('4.2', 'Backend Implementation', ''),
    ('4.3', 'AI Engine Implementation', ''),
    ('4.4', 'Screenshots', ''),
    ('4.5', 'Code Samples', ''),
    ('5.', 'Testing', ''),
    ('5.1', 'Test Cases', ''),
    ('5.2', 'Test Results', ''),
    ('6.', 'Deployment', ''),
    ('6.1', 'Cloud Deployment', ''),
    ('6.2', 'System Requirements', ''),
    ('7.', 'Conclusion', ''),
    ('7.1', 'Summary', ''),
    ('7.2', 'Future Enhancements', ''),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
    run = p.add_run(f'{num}\t{title}')
    if '.' not in num or num.endswith('.'):
        run.bold = True

doc.add_page_break()

# ==========================================
# 1. PROJECT INTRODUCTION
# ==========================================
add_heading('1. Project Introduction', 1)

add_heading('1.1 Project Overview', 2)
doc.add_paragraph(
    'The AI Personality & Career Recommendation System is an intelligent web-based application '
    'that analyzes user personality using Natural Language Processing (NLP) and the Big Five Personality Model. '
    'The system helps users discover their personality traits, identifies suitable career paths, '
    'and connects them with similar-minded individuals who share comparable goals and interests.'
)

doc.add_paragraph(
    'The project combines multiple modern technologies including Flask-based backend, MongoDB database, '
    'React-ready frontend architecture, and sophisticated AI algorithms for personality analysis and recommendation.'
)

add_heading('1.2 Project Name', 2)
add_bullet('English: AI Personality & Career Recommendation System')
add_bullet('Arabic: نظام ذكي لتحليل شخصية المستخدم وتوصية المسارات المهنية والأصدقاء')
add_bullet('Short Name: PersonaAI')

add_heading('1.3 Project Objectives', 2)
objectives = [
    'Create a user-friendly system for personality analysis based on the Big Five Personality Model',
    'Provide personalized career recommendations based on personality traits and user goals',
    'Enable users to discover and connect with similar-minded individuals',
    'Implement text analysis capabilities for personality extraction from user-written content',
    'Develop a scalable architecture that can handle multiple users simultaneously',
    'Provide bilingual support (English and Arabic) for broader accessibility',
    'Generate comprehensive personality reports with visual representations',
]
for obj in objectives:
    add_bullet(obj)

add_heading('1.4 Project Scope', 2)
doc.add_paragraph('The system includes the following core features:')

scope_items = [
    ('User Registration & Authentication', 'Users can create accounts, login, and manage their profiles securely'),
    ('Personality Quiz', 'A comprehensive 50-question quiz based on the Big Five personality model'),
    ('Text-Based Analysis', 'NLP-powered analysis of free-text descriptions provided by users'),
    ('Career Recommendations', 'Personalized career suggestions based on personality trait analysis'),
    ('User Similarity Matching', 'Finding and recommending similar users based on personality compatibility'),
    ('Responsive UI', 'Modern, responsive web interface accessible on all devices'),
    ('Bilingual Support', 'Full English and Arabic language support'),
]
for title, desc in scope_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==========================================
# 2. SYSTEM ANALYSIS
# ==========================================
add_heading('2. System Analysis', 1)

add_heading('2.1 Requirements Analysis', 2)

add_heading('2.1.1 Functional Requirements', 3)
func_reqs = [
    ('FR-001', 'User Registration', 'System shall allow new users to register with name, email, password, age, and career goal'),
    ('FR-002', 'User Authentication', 'System shall authenticate users using email and password'),
    ('FR-003', 'Personality Quiz', 'System shall present 50 questions covering the Big Five personality traits'),
    ('FR-004', 'Quiz Answer Storage', 'System shall store user quiz answers for future reference'),
    ('FR-005', 'Personality Analysis', 'System shall calculate five personality trait scores (0-100%)'),
    ('FR-006', 'Personality Type Classification', 'System shall classify users into personality types based on dominant traits'),
    ('FR-007', 'Career Recommendations', 'System shall recommend careers based on personality analysis'),
    ('FR-008', 'Skill Suggestions', 'System shall suggest skills needed for recommended careers'),
    ('FR-009', 'Similar User Discovery', 'System shall find users with similar personality profiles'),
    ('FR-010', 'Text Analysis', 'System shall analyze free-text input to extract personality traits'),
    ('FR-011', 'Profile Management', 'System shall allow users to view and manage their profiles'),
    ('FR-012', 'Bilingual Support', 'System shall support English and Arabic languages'),
]
add_table_with_headers(['ID', 'Requirement', 'Description'], func_reqs)

doc.add_paragraph()
add_heading('2.1.2 Non-Functional Requirements', 3)
non_func_reqs = [
    ('NFR-001', 'Performance', 'System shall respond to user requests within 3 seconds under normal load'),
    ('NFR-002', 'Scalability', 'System architecture shall support horizontal scaling for increased users'),
    ('NFR-003', 'Security', 'User passwords shall be hashed using bcrypt before storage'),
    ('NFR-004', 'Usability', 'Interface shall be intuitive and require minimal training'),
    ('NFR-005', 'Availability', 'System shall be available 99% of the time'),
    ('NFR-006', 'Data Integrity', 'User data shall be protected from corruption and loss'),
    ('NFR-007', 'Accessibility', 'System shall be accessible on desktop, tablet, and mobile devices'),
    ('NFR-008', 'Maintainability', 'Code shall be well-documented and follow coding standards'),
]
add_table_with_headers(['ID', 'Quality Attribute', 'Description'], non_func_reqs)

doc.add_paragraph()

add_heading('2.2 System Context Diagram', 3)
doc.add_paragraph(
    'The context diagram illustrates the system boundaries and external entities interacting with the AI Personality & Career Recommendation System.'
)
doc.add_paragraph()

# Context diagram description
context_p = doc.add_paragraph()
context_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
context_p.add_run('┌─────────────────────────────────────────────────────────────────┐\n').bold = True
context_p.add_run('│           AI Personality & Career Recommendation System          │\n')
context_p.add_run('│                                                                 │\n')
context_p.add_run('│  ┌─────────────┐  ┌─────────────┐  ┌─────────────────────────┐  │\n')
context_p.add_run('│  │  Frontend   │  │  Backend    │  │   AI Engine & Database  │  │\n')
context_p.add_run('│  │  (HTML/CSS) │  │  (Flask)    │  │   (NLP + MongoDB)       │  │\n')
context_p.add_run('│  └─────────────┘  └─────────────┘  └─────────────────────────┘  │\n')
context_p.add_run('└─────────────────────────────────────────────────────────────────┘\n')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('External Entities:').bold = True
doc.add_paragraph()

ext_entities = [
    'User: Interacts with the system through web interface',
    'MongoDB: Stores user data, personality results, and career profiles',
    'Browser: Renders the frontend interface',
]
for entity in ext_entities:
    add_bullet(entity)

add_heading('2.3 Use Case Diagram', 3)
doc.add_paragraph('The use case diagram shows the interactions between users and the system:')
doc.add_paragraph()

uc_desc = '''┌─────────────────────────────────────────────────────────────┐
│                    AI Personality System                           │
│                                                             │

│  ┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐  │
│  │ Register │    │  Login   │    │  Take    │    │  View    │  │
│  │          │    │          │    │  Quiz    │    │  Results │  │
│  └────┬─────┘    └────┬─────┘    └────┬─────┘    └────┬─────┘  │
│       │               │               │               │        │
│       └───────────────┴───────────────┴───────────────┘        │
│                            │                                    │
│                            ▼                                    │
│  ┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐  │
│  │ View     │    │ View     │    │ Connect  │    │  Text    │  │
│  │ Profile  │    │ Careers  │    │ with     │    │  Analysis│  │
│  │          │    │          │    │ Users    │    │          │  │
│  └──────────┘    └──────────┘    └──────────┘    └──────────┘  │
└─────────────────────────────────────────────────────────────┘'''
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(uc_desc)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading('2.4 Use Case Descriptions', 3)

use_cases = [
    {
        'id': 'UC-001',
        'name': 'User Registration',
        'actor': 'Guest User',
        'description': 'New user creates an account by providing name, email, password, age, and career goal',
        'preconditions': 'User is not logged in',
        'flow': [
            '1. User clicks "Register" button',
            '2. System displays registration form',
            '3. User fills in all required fields',
            '4. User submits the form',
            '5. System validates input and creates account',
            '6. System redirects to personality quiz'
        ],
        'postconditions': 'User account is created and user is logged in'
    },
    {
        'id': 'UC-002',
        'name': 'User Login',
        'actor': 'Registered User',
        'description': 'Registered user logs into the system',
        'preconditions': 'User has an existing account',
        'flow': [
            '1. User enters email and password',
            '2. System validates credentials',
            '3. System authenticates user',
            '4. System redirects to quiz (if not taken) or profile (if taken)'
        ],
        'postconditions': 'User is authenticated and session is created'
    },
    {
        'id': 'UC-003',
        'name': 'Take Personality Quiz',
        'actor': 'Authenticated User',
        'description': 'User answers 50 personality questions',
        'preconditions': 'User is logged in and has not taken the quiz',
        'flow': [
            '1. User navigates to quiz page',
            '2. System displays first question',
            '3. User selects answer (1-5 scale)',
            '4. User navigates to next question',
            '5. Repeat steps 3-4 for all questions',
            '6. User submits quiz',
            '7. System calculates personality traits',
            '8. System displays results'
        ],
        'postconditions': 'User personality profile is generated and stored'
    },
    {
        'id': 'UC-004',
        'name': 'View Personality Results',
        'actor': 'Authenticated User',
        'description': 'User views their personality analysis and recommendations',
        'preconditions': 'User has completed the personality quiz',
        'flow': [
            '1. User navigates to results/profile page',
            '2. System retrieves personality data',
            '3. System displays personality traits with visual charts',
            '4. System displays career recommendations',
            '5. System displays similar users'
        ],
        'postconditions': 'User can view their complete personality profile'
    },
    {
        'id': 'UC-005',
        'name': 'Connect with Similar Users',
        'actor': 'Authenticated User',
        'description': 'User views and connects with similar personality profiles',
        'preconditions': 'User has completed quiz and other users exist',
        'flow': [
            '1. User views recommended similar users',
            '2. User clicks on a user profile',
            '3. System displays user details',
            '4. User can view similarity score and traits comparison'
        ],
        'postconditions': 'User can view other users with similar personalities'
    },
    {
        'id': 'UC-006',
        'name': 'Text-Based Personality Analysis',
        'actor': 'Authenticated User',
        'description': 'User provides free-text for personality analysis',
        'preconditions': 'User is logged in',
        'flow': [
            '1. User enters text about themselves',
            '2. System processes text using NLP',
            '3. System extracts keywords and patterns',
            '4. System calculates personality traits',
            '5. System displays analysis results'
        ],
        'postconditions': 'Personality analysis based on text is generated'
    },
]

for uc in use_cases:
    p = doc.add_paragraph()
    p.add_run(f"Use Case: {uc['name']} ({uc['id']})").bold = True
    doc.add_paragraph(f"Actor: {uc['actor']}")
    doc.add_paragraph(f"Description: {uc['description']}")
    
    p = doc.add_paragraph()
    p.add_run('Preconditions: ').bold = True
    p.add_run(uc['preconditions'])
    
    p = doc.add_paragraph()
    p.add_run('Flow of Events:').bold = True
    for step in uc['flow']:
        add_bullet(step, level=1)
    
    p = doc.add_paragraph()
    p.add_run('Postconditions: ').bold = True
    p.add_run(uc['postconditions'])
    doc.add_paragraph()

add_heading('2.5 Wireframes', 3)
doc.add_paragraph(
    'The following wireframes illustrate the key user interface screens of the system:'
)

wireframes = [
    ('Home Page', 'Landing page with hero section, features overview, and call-to-action'),
    ('Login/Register Page', 'Authentication form with tabbed interface for login and registration'),
    ('Personality Quiz', 'Question card with 1-5 rating buttons and progress indicator'),
    ('Results Page', 'Personality traits visualization with career recommendations'),
    ('User Profile', 'Complete user profile with traits, careers, and similar users'),
]

for name, desc in wireframes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph('Note: Detailed wireframes are provided in the accompanying design documentation.')

doc.add_page_break()

# ==========================================
# 3. SYSTEM DESIGN
# ==========================================
add_heading('3. System Design', 1)

add_heading('3.1 Database Design', 2)
doc.add_paragraph(
    'The system uses MongoDB as its primary database. MongoDB was chosen for its flexibility '
    'with document storage, easy scalability, and excellent JSON-like query support.'
)

add_heading('3.2 Entity Relationship Diagram (ERD)', 3)

# ERD description
erd_diagram = '''
┌─────────────────┐         ┌─────────────────┐
│      User       │         │   Personality   │
├─────────────────┤         ├─────────────────┤
│ _id (PK)        │──┐      │ _id (PK)        │
│ name            │  │      │ user_email (FK)│◄─┘
│ email (UNIQUE)  │  └──────│ openness        │
│ password (hash) │         │ conscientious. │
│ age             │         │ extraversion    │
│ goal            │         │ agreeableness  │
│ created_at      │         │ neuroticism    │
└─────────────────┘         │ personality_type│
       │                    │ answers[]       │
       │                    └─────────────────┘
       │
       ▼
┌─────────────────┐         ┌─────────────────┐
│   Career        │         │  Recommendation │
├─────────────────┤         ├─────────────────┤
│ _id (PK)        │         │ _id (PK)        │
│ name            │         │ user_email (FK) │
│ description     │         │ recommended_    │
│ skills[]        │         │   careers[]     │
│ salary_range    │         │ similar_users[] │
│ traits_profile  │         └─────────────────┘
└─────────────────┘
'''
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(erd_diagram)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading('3.3 Database Mapping', 3)

add_heading('User Collection', 4)
user_schema = [
    ('_id', 'ObjectId', 'Primary key, auto-generated'),
    ('name', 'String', 'User full name, required'),
    ('email', 'String', 'Unique email address, indexed'),
    ('password', 'String', 'Bcrypt hashed password'),
    ('age', 'Number', 'User age (optional)'),
    ('goal', 'String', 'Career goal description'),
    ('created_at', 'DateTime', 'Account creation timestamp'),
]
add_table_with_headers(['Field', 'Type', 'Description'], user_schema)

doc.add_paragraph()
add_heading('Personality Collection', 4)
personality_schema = [
    ('_id', 'ObjectId', 'Primary key, auto-generated'),
    ('user_email', 'String', 'Reference to user email, indexed'),
    ('user_id', 'String', 'Reference to user _id'),
    ('openness', 'Number', 'Openness score (0-100)'),
    ('conscientiousness', 'Number', 'Conscientiousness score (0-100)'),
    ('extraversion', 'Number', 'Extraversion score (0-100)'),
    ('agreeableness', 'Number', 'Agreeableness score (0-100)'),
    ('neuroticism', 'Number', 'Neuroticism score (0-100)'),
    ('personality_type', 'String', 'Derived personality type'),
    ('answers', 'Array', 'Array of 50 quiz answers'),
    ('created_at', 'DateTime', 'Analysis timestamp'),
]
add_table_with_headers(['Field', 'Type', 'Description'], personality_schema)

add_heading('3.4 Service Architecture', 3)
doc.add_paragraph(
    'The system follows a microservices-inspired architecture with clear separation of concerns:'
)

services = [
    ('Frontend Service', 'HTML5, CSS3, JavaScript', 'User interface, responsive design, bilingual support'),
    ('API Gateway', 'Flask', 'Route handling, CORS management, request validation'),
    ('Authentication Service', 'Flask + bcrypt', 'User registration, login, session management'),
    ('Personality Analysis Service', 'Python + NLP', 'Trait calculation, personality classification'),
    ('Recommendation Engine', 'Python + NumPy', 'Career matching, user similarity using cosine similarity'),
    ('Database Service', 'MongoDB', 'Data persistence, indexing, queries'),
]
add_table_with_headers(['Service', 'Technology', 'Responsibility'], services)

doc.add_page_break()

# ==========================================
# 4. IMPLEMENTATION
# ==========================================
add_heading('4. Implementation', 1)

add_heading('4.1 Frontend Implementation', 2)
doc.add_paragraph(
    'The frontend is built using modern HTML5, CSS3, and vanilla JavaScript with no heavy frameworks, '
    'ensuring fast load times and easy deployment.'
)

frontend_features = [
    'Responsive design with CSS Grid and Flexbox',
    'Modern glassmorphism UI with backdrop filters',
    'Real-time quiz progress tracking',
    'Interactive personality trait visualizations',
    'Bilingual support (EN/AR) with RTL layout',
    'Local storage for offline quiz capability',
    'Smooth animations and transitions',
]
for feature in frontend_features:
    add_bullet(feature)

add_heading('4.2 Backend Implementation', 2)
doc.add_paragraph(
    'The backend is built with Flask, providing a lightweight yet powerful API framework.'
)

backend_components = [
    ('app.py', 'Main Flask application with route registration'),
    ('/backend/routes/auth.py', 'Authentication endpoints (register, login, user profile)'),
    ('/backend/routes/personality.py', 'Personality analysis and storage endpoints'),
    ('/backend/routes/recommendations.py', 'Career and user recommendation endpoints'),
    ('/backend/config.py', 'Configuration management'),
]

for path, desc in backend_components:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{path}: ')
    run.bold = True
    run.font.name = 'Courier New'
    p.add_run(desc)

add_heading('4.3 AI Engine Implementation', 2)
doc.add_paragraph(
    'The AI engine handles personality analysis and recommendation generation.'
)

ai_components = [
    ('PersonalityAnalyzer', 'Calculates Big Five trait scores from quiz answers and text'),
    ('TextProcessor', 'Preprocesses text using NLP techniques (stopword removal, tokenization)'),
    ('CareerRecommender', 'Matches user traits to career profiles using weighted scoring'),
    ('UserRecommender', 'Finds similar users using cosine similarity algorithm'),
    ('Similarity Module', 'Implements cosine, euclidean, and manhattan distance calculations'),
]

for name, desc in ai_components:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

add_heading('4.4 Screenshots', 3)
doc.add_paragraph('The following screenshots demonstrate the key features of the system:')

screenshots = [
    ('Figure 1: Home Page', 'Landing page with hero section and feature highlights'),
    ('Figure 2: Login/Register', 'Authentication interface with bilingual support'),
    ('Figure 3: Personality Quiz', 'Interactive quiz with progress bar and answer selection'),
    ('Figure 4: Results Dashboard', 'Personality traits visualization with percentage scores'),
    ('Figure 5: Career Recommendations', 'Matched careers with skill requirements'),
    ('Figure 6: User Profile', 'Complete profile with similar users section'),
]

for name, desc in screenshots:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph('Note: Screenshots are provided in the accompanying visual documentation.')

add_heading('4.5 Code Samples', 3)
doc.add_paragraph('Key implementation code snippets:')

# Code sample 1
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Personality Analysis (ai_engine/personality_analyzer.py):').bold = True

code1 = '''
class PersonalityAnalyzer:
    TRAIT_NAMES = ['openness', 'conscientiousness', 'extraversion', 
                    'agreeableness', 'neuroticism']
    
    def calculate_traits(self, answers):
        if len(answers) != 50:
            raise ValueError("Expected 50 answers")
        
        openness = np.mean(answers[0:10]) * 20
        conscientiousness = np.mean(answers[10:20]) * 20
        extraversion = np.mean(answers[20:30]) * 20
        agreeableness = np.mean(answers[30:40]) * 20
        neuroticism = np.mean(answers[40:50]) * 20
        
        return {
            'openness': round(openness, 2),
            'conscientiousness': round(conscientiousness, 2),
            'extraversion': round(extraversion, 2),
            'agreeableness': round(agreeableness, 2),
            'neuroticism': round(neuroticism, 2)
        }
'''
p = doc.add_paragraph()
run = p.add_run(code1)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Code sample 2
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Cosine Similarity (ai_engine/recommender/similarity.py):').bold = True

code2 = '''
def cosine_similarity(vec1, vec2):
    if not vec1 or not vec2:
        return 0.0
    
    vec1 = np.array(vec1)
    vec2 = np.array(vec2)
    
    dot_product = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)
    
    if norm1 == 0 or norm2 == 0:
        return 0.0
    
    return dot_product / (norm1 * norm2)
'''
p = doc.add_paragraph()
run = p.add_run(code2)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ==========================================
# 5. TESTING
# ==========================================
add_heading('5. Testing', 1)

add_heading('5.1 Test Cases', 2)
doc.add_paragraph(
    'Comprehensive testing was conducted to ensure system reliability and correctness.'
)

test_cases = [
    {
        'id': 'TC-001',
        'category': 'Authentication',
        'description': 'User registration with valid data',
        'steps': ['1. Navigate to register page', '2. Fill all required fields', '3. Submit form'],
        'expected': 'Account created, redirected to quiz',
        'status': 'PASS'
    },
    {
        'id': 'TC-002',
        'category': 'Authentication',
        'description': 'User registration with duplicate email',
        'steps': ['1. Attempt to register with existing email'],
        'expected': 'Error message displayed',
        'status': 'PASS'
    },
    {
        'id': 'TC-003',
        'category': 'Authentication',
        'description': 'User login with correct credentials',
        'steps': ['1. Enter email and password', '2. Submit login form'],
        'expected': 'Login successful, redirected to profile/quiz',
        'status': 'PASS'
    },
    {
        'id': 'TC-004',
        'category': 'Personality Quiz',
        'description': 'Complete personality quiz',
        'steps': ['1. Answer all 50 questions', '2. Submit quiz'],
        'expected': 'Personality results displayed',
        'status': 'PASS'
    },
    {
        'id': 'TC-005',
        'category': 'Personality Quiz',
        'description': 'Quiz navigation (next/previous)',
        'steps': ['1. Answer question', '2. Navigate to next', '3. Navigate back'],
        'expected': 'Answers preserved during navigation',
        'status': 'PASS'
    },
    {
        'id': 'TC-006',
        'category': 'Analysis',
        'description': 'Personality trait calculation',
        'steps': ['1. Complete quiz', '2. View results'],
        'expected': 'Correct trait percentages (0-100)',
        'status': 'PASS'
    },
    {
        'id': 'TC-007',
        'category': 'Recommendation',
        'description': 'Career recommendations generated',
        'steps': ['1. View personality results'],
        'expected': '4-5 career matches with percentages',
        'status': 'PASS'
    },
    {
        'id': 'TC-008',
        'category': 'Recommendation',
        'description': 'Similar users discovery',
        'steps': ['1. View profile page'],
        'expected': 'List of similar users displayed',
        'status': 'PASS'
    },
    {
        'id': 'TC-009',
        'category': 'UI/UX',
        'description': 'Language switching (EN/AR)',
        'steps': ['1. Click language toggle'],
        'expected': 'Interface switches to Arabic with RTL',
        'status': 'PASS'
    },
    {
        'id': 'TC-010',
        'category': 'UI/UX',
        'description': 'Responsive design',
        'steps': ['1. View site on mobile/tablet'],
        'expected': 'Layout adjusts properly',
        'status': 'PASS'
    },
]

add_table_with_headers(
    ['ID', 'Category', 'Description', 'Expected Result', 'Status'],
    [[tc['id'], tc['category'], tc['description'], tc['expected'], tc['status']] for tc in test_cases]
)

add_heading('5.2 Test Results', 3)
doc.add_paragraph('Testing Summary:')

test_summary = [
    ('Total Test Cases', '10'),
    ('Passed', '10'),
    ('Failed', '0'),
    ('Pass Rate', '100%'),
]
for label, value in test_summary:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(value)

doc.add_paragraph()
doc.add_paragraph(
    'All critical functionality has been tested and verified. The system handles edge cases '
    'such as invalid input, duplicate emails, and missing data gracefully with appropriate error messages.'
)

doc.add_page_break()

# ==========================================
# 6. DEPLOYMENT
# ==========================================
add_heading('6. Deployment', 1)

add_heading('6.1 Cloud Deployment', 2)
doc.add_paragraph(
    'The system can be deployed on various cloud platforms. Below is the recommended deployment architecture:'
)

deployment_options = [
    ('Frontend', 'Static hosting (Vercel, Netlify, GitHub Pages)', 'No server required, CDN-backed delivery'),
    ('Backend', 'Python hosting (Railway, Render, Heroku)', 'Flask application with gunicorn'),
    ('Database', 'MongoDB Atlas', 'Managed MongoDB with automatic scaling'),
]

add_table_with_headers(['Component', 'Recommended Platform', 'Notes'], deployment_options)

add_heading('Deployment Steps', 3)
deploy_steps = [
    '1. Clone the repository to local development environment',
    '2. Set up MongoDB Atlas cluster and obtain connection string',
    '3. Configure environment variables (MONGO_URI, SECRET_KEY)',
    '4. Deploy backend to chosen platform (e.g., Railway/Render)',
    '5. Configure CORS origins in backend configuration',
    '6. Deploy frontend to static hosting (e.g., Vercel)',
    '7. Update frontend API endpoints to production backend URL',
    '8. Test all functionality in production environment',
]
for step in deploy_steps:
    add_bullet(step)

add_heading('6.2 System Requirements', 3)

add_heading('Client Requirements', 4)
client_reqs = [
    'Modern web browser (Chrome, Firefox, Safari, Edge)',
    'JavaScript enabled',
    'Minimum screen resolution: 320px width',
    'Stable internet connection',
]
for req in client_reqs:
    add_bullet(req)

add_heading('Server Requirements', 4)
server_reqs = [
    'Python 3.8+',
    'MongoDB 4.4+ or MongoDB Atlas',
    'RAM: 512MB minimum (1GB recommended)',
    'Storage: 1GB for application and logs',
    'Network: HTTPS enabled for production',
]
for req in server_reqs:
    add_bullet(req)

doc.add_page_break()

# ==========================================
# 7. CONCLUSION
# ==========================================
add_heading('7. Conclusion', 1)

add_heading('7.1 Summary', 2)
doc.add_paragraph(
    'The AI Personality & Career Recommendation System successfully implements a comprehensive '
    'personality analysis solution using the Big Five Personality Model. The system provides '
    'valuable insights into user personality traits and offers personalized career recommendations '
    'and social connections.'
)

doc.add_paragraph('Key achievements of the project include:')

achievements = [
    'Implemented a complete personality analysis system using scientific Big Five model',
    'Developed an intuitive, responsive user interface with bilingual support',
    'Created sophisticated recommendation algorithms for careers and user matching',
    'Built a scalable Flask backend with RESTful API architecture',
    'Integrated NLP capabilities for text-based personality analysis',
    'Achieved 100% test pass rate with comprehensive test coverage',
    'Designed system for easy cloud deployment and scaling',
]
for achievement in achievements:
    add_bullet(achievement)

add_heading('7.2 Future Enhancements', 2)
doc.add_paragraph(
    'The following enhancements are planned for future development:'
)

future_enhancements = [
    ('Machine Learning Models', 'Train custom ML models for more accurate personality prediction'),
    ('Mobile Application', 'Develop native iOS and Android apps for better mobile experience'),
    ('Social Features', 'Add messaging, networking, and collaboration features'),
    ('Career Database Expansion', 'Expand career profiles with more details and pathways'),
    ('Assessment Tools', 'Add skills assessments and aptitude tests'),
    ('Progress Tracking', 'Track personality changes over time with periodic reassessments'),
    ('Integration APIs', 'Provide APIs for third-party integration'),
    ('AI Chatbot', 'Add conversational AI for personalized career guidance'),
    ('Data Analytics Dashboard', 'Admin dashboard for system analytics and insights'),
    ('Multi-language Support', 'Add more languages beyond English and Arabic'),
]

for enhancement, desc in future_enhancements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{enhancement}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

# Final note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('--- End of Document ---').italic = True

# Save the document
doc.save('/mnt/c/Users/BAN/testme/AI_Personality_Career_Recommendation_System_Documentation.docx')
print('Document created successfully!')
